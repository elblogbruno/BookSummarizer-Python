from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer,LTChar,LTImage,LTTextBoxHorizontal
# Import summarizer
from summarizer import Summarizer
from docx import Document
from docx.shared import Inches


class BookSummarizer():
    def __init__(self, pdf_path,pages,doc_name,save_url):
        self.model = Summarizer()
        self.document = Document()

        self.pages = pages
        self.doc_name = doc_name

        self.save_url = save_url
        self.path = pdf_path

    def summarizeBook(self):
        self.document.add_heading(self.doc_name, 0)
        for page_layout in extract_pages(self.path,page_numbers=self.pages):
            
            for element in page_layout:
                print(element)
                if isinstance(element, LTTextContainer):
                    text = element.get_text()
                    if text.strip():
                        for c in  o._objs:
                            if isinstance(c, pdfminer.layout.LTChar):
                                print("fontname %s"%c.fontname)
                    if self.is_header(text) == False:
                        result = self.model(text, ratio=0.2)
                        p = self.document.add_paragraph(result)
                    else:
                        p = self.document.add_heading(text, level=1)
                        
        self.document.save(self.save_url+self.doc_name)

    def is_header(self,text):
        for character in text:
            if character.isdigit():
                return True
                break
        return False

